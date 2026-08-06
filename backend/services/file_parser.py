"""文件解析器 - 支持 docx/pdf/txt/xlsx/xls/csv/doc
V6.17: 新增 xlsx/xls/csv 支持，旧格式 .doc 基础支持

纯本地解析，不调外部 API
"""
import os
import logging

logger = logging.getLogger(__name__)

MAX_CHARS = 50000  # 超长文档截断阈值


def parse(file_path: str) -> str:
    """
    解析文件内容为纯文本
    :param file_path: 文件路径
    :return: 解析后的文本
    :raises: ValueError 不支持的文件类型
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.docx':
        text = _parse_docx(file_path)
    elif ext == '.pdf':
        text = _parse_pdf(file_path)
    elif ext == '.txt':
        text = _parse_txt(file_path)
    elif ext == '.csv':
        text = _parse_csv(file_path)
    elif ext in ('.xlsx', '.xls'):
        text = _parse_xlsx(file_path)
    elif ext == '.doc':
        text = _parse_legacy_doc(file_path)
    else:
        raise ValueError(f'不支持的文件类型: {ext}，仅支持 .docx/.pdf/.txt/.xlsx/.xls/.csv/.doc')

    # 超长截断
    if len(text) > MAX_CHARS:
        logger.warning(f"文档超长({len(text)}字)，截断到 {MAX_CHARS} 字: {file_path}")
        text = text[:MAX_CHARS]

    return text


def _parse_docx(file_path: str) -> str:
    """python-docx 解析 .docx"""
    from docx import Document
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return '\n'.join(paragraphs)


def _parse_pdf(file_path: str) -> str:
    """pdfplumber 解析 .pdf"""
    import pdfplumber
    texts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
            # 尝试提取表格
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    for row in table:
                        if row:
                            row_text = ' | '.join(str(c or '').strip() for c in row if c)
                            if row_text:
                                texts.append(row_text)
    return '\n'.join(texts)


def _parse_txt(file_path: str) -> str:
    """txt 解析，自动探测 utf-8/gbk"""
    # 尝试 utf-8
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        pass
    # 回退 gbk
    try:
        with open(file_path, 'r', encoding='gbk') as f:
            return f.read()
    except UnicodeDecodeError:
        pass
    # 最后尝试 latin-1（不会失败）
    with open(file_path, 'r', encoding='latin-1') as f:
        return f.read()


def _parse_csv(file_path: str) -> str:
    """CSV 文件解析 - 转为可读文本表格"""
    import csv
    
    # 尝试不同编码
    for encoding in ['utf-8-sig', 'utf-8', 'gbk', 'gb2312', 'latin-1']:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                # 检测分隔符
                sample = f.read(2048)
                f.seek(0)
                try:
                    dialect = csv.Sniffer().sniff(sample, delimiters=',;\t|')
                except csv.Error:
                    dialect = csv.excel
                
                reader = csv.reader(f, dialect)
                rows = list(reader)
                
                if not rows:
                    return ''
                
                # 计算列宽用于对齐
                col_widths = []
                for col_idx in range(max(len(r) for r in rows[:20]) if rows else 0):
                    max_w = 0
                    for row in rows[:50]:  # 只看前50行计算宽度
                        if col_idx < len(row):
                            max_w = max(max_w, len(str(row[col_idx])))
                    col_widths.append(min(max_w, 30))  # 最大列宽30
                
                # 转为文本格式
                lines = []
                for row_idx, row in enumerate(rows[:200]):  # 限制行数
                    cells = []
                    for col_idx, cell in enumerate(row):
                        w = col_widths[col_idx] if col_idx < len(col_widths) else 10
                        cells.append(str(cell).ljust(w)[:w])
                    lines.append(' | '.join(cells))
                    if row_idx == 0:
                        lines.append('-+-'.join('-' * w for w in col_widths))
                
                return '\n'.join(lines)
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.warning(f"CSV解析失败({encoding}): {e}")
            continue
    
    return ''


def _parse_xlsx(file_path: str) -> str:
    """Excel 文件解析 - 使用 openpyxl (xlsx) 或 xlrd (xls)"""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.xlsx':
        return _parse_xlsx_openpyxl(file_path)
    elif ext == '.xls':
        return _parse_xls_fallback(file_path)
    return ''


def _parse_xlsx_openpyxl(file_path: str) -> str:
    """使用 openpyxl 解析 .xlsx 文件"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        lines = []
        
        for sheet_name in wb.sheetnames[:5]:  # 最多5个sheet
            ws = wb[sheet_name]
            if len(wb.sheetnames) > 1:
                lines.append(f'\n=== 工作表: {sheet_name} ===')
            
            rows_data = []
            for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
                if row_idx >= 500:  # 限制行数
                    lines.append(f'... (还有更多行)')
                    break
                cells = [str(c) if c is not None else '' for c in row]
                if any(c for c in cells):  # 跳过全空行
                    rows_data.append(cells)
            
            if rows_data:
                # 计算列宽
                max_cols = max(len(r) for r in rows_data)
                col_widths = []
                for col_idx in range(max_cols):
                    max_w = 0
                    for row in rows_data[:50]:
                        if col_idx < len(row):
                            max_w = max(max_w, len(row[col_idx]))
                    col_widths.append(min(max_w, 25))
                
                # 格式化输出
                for row_idx, row in enumerate(rows_data):
                    cells = []
                    for col_idx, cell in enumerate(row):
                        w = col_widths[col_idx] if col_idx < len(col_widths) else 10
                        cells.append(cell.ljust(w)[:w])
                    lines.append(' | '.join(cells))
                    if row_idx == 0:
                        lines.append('-+-'.join('-' * w for w in col_widths))
        
        wb.close()
        return '\n'.join(lines)
    except Exception as e:
        logger.error(f"openpyxl 解析失败: {e}")
        return f'[Excel文件解析失败: {str(e)}]'


def _parse_xls_fallback(file_path: str) -> str:
    """尝试解析旧版 .xls 文件"""
    try:
        import pandas as pd
        df_dict = pd.read_excel(file_path, sheet_name=None, nrows=500)
        lines = []
        for sheet_name, df in list(df_dict.items())[:5]:
            if len(df_dict) > 1:
                lines.append(f'\n=== 工作表: {sheet_name} ===')
            lines.append(df.to_string(index=False, max_rows=200))
        return '\n'.join(lines)
    except Exception as e:
        logger.warning(f"pandas xls解析失败: {e}")
        return f'[Excel文件解析失败: {str(e)}，建议转换为 .xlsx 格式]'


def _parse_legacy_doc(file_path: str) -> str:
    """尝试解析旧版 .doc 文件 - 有限支持"""
    # 尝试用 antiword 或 textract（如果安装了）
    # 否则返回提示信息
    try:
        # 尝试用 subprocess 调用 antiword (如果可用)
        import subprocess
        result = subprocess.run(
            ['antiword', file_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (FileNotFoundError, subprocess.SubprocessError):
        pass
    
    # 尝试用 catdoc (如果可用)
    try:
        import subprocess
        result = subprocess.run(
            ['catdoc', file_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (FileNotFoundError, subprocess.SubprocessError):
        pass
    
    # 尝试提取可读文本（简单二进制提取）
    try:
        with open(file_path, 'rb') as f:
            data = f.read()
        # 提取 ASCII 和 UTF-8 可打印字符
        import re
        text = data.decode('utf-8', errors='ignore')
        # 提取有意义的文本片段
        fragments = re.findall(r'[\u4e00-\u9fff\w\s.,;:!?，。；：！？、（）()]+', text)
        meaningful = [f for f in fragments if len(f.strip()) > 5]
        if meaningful:
            return '\n'.join(meaningful[:100])
    except Exception:
        pass
    
    return '[.doc 格式需要转换为 .docx 以获得完整文本提取。建议用 Word 另存为 .docx 格式后重新上传。]'
