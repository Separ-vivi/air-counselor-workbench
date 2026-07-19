"""知识库 + FAQ + 文书生成 + 周汇总 路由"""
import io
import json
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from typing import Optional
from database import get_db
from models import KnowledgeDoc, FAQ, DocumentTemplate, GeneratedDocument, Student, WeeklySummary

router = APIRouter(prefix='/api')


# ===== 知识库 =====
@router.get('/knowledge')
def list_knowledge(db: Session = Depends(get_db)):
    items = db.query(KnowledgeDoc).order_by(KnowledgeDoc.created_at.desc()).all()
    return [{
        'id': d.id, 'title': d.title, 'doc_type': d.doc_type,
        'content': d.content[:200] + '...' if len(d.content) > 200 else d.content,
        'chunk_count': d.chunk_count, 'created_at': str(d.created_at),
    } for d in items]


@router.get('/knowledge/{kid}')
def get_knowledge(kid: int, db: Session = Depends(get_db)):
    d = db.query(KnowledgeDoc).get(kid)
    if not d:
        raise HTTPException(404)
    return {'id': d.id, 'title': d.title, 'content': d.content, 'doc_type': d.doc_type}


@router.post('/knowledge')
def create_knowledge(data: dict, db: Session = Depends(get_db)):
    d = KnowledgeDoc(**data)
    db.add(d)
    db.commit()
    db.refresh(d)
    return {'id': d.id}


@router.delete('/knowledge/{kid}')
def delete_knowledge(kid: int, db: Session = Depends(get_db)):
    d = db.query(KnowledgeDoc).get(kid)
    if d:
        db.delete(d)
        db.commit()
    return {'ok': True}


# ===== FAQ =====
@router.get('/faqs')
def list_faqs(published_only: bool = False, db: Session = Depends(get_db)):
    q = db.query(FAQ)
    if published_only:
        q = q.filter(FAQ.is_published == True)
    items = q.order_by(FAQ.created_at.desc()).all()
    return [{
        'id': f.id, 'question': f.question, 'answer': f.answer,
        'category': f.category, 'is_published': f.is_published,
    } for f in items]


@router.get('/faqs/{faq_id}')
def get_faq(faq_id: int, db: Session = Depends(get_db)):
    f = db.query(FAQ).get(faq_id)
    if not f:
        raise HTTPException(404, 'FAQ不存在')
    return {
        'id': f.id, 'question': f.question, 'answer': f.answer,
        'category': f.category, 'is_published': f.is_published,
    }


@router.post('/faqs')
def create_faq(data: dict, db: Session = Depends(get_db)):
    f = FAQ(**data)
    db.add(f)
    db.commit()
    db.refresh(f)
    return {'id': f.id}


@router.put('/faqs/{fid}')
def update_faq(fid: int, data: dict, db: Session = Depends(get_db)):
    f = db.query(FAQ).get(fid)
    if not f:
        raise HTTPException(404)
    for k, v in data.items():
        if hasattr(f, k):
            setattr(f, k, v)
    db.commit()
    return {'ok': True}


@router.delete('/faqs/{fid}')
def delete_faq(fid: int, db: Session = Depends(get_db)):
    f = db.query(FAQ).get(fid)
    if f:
        db.delete(f)
        db.commit()
    return {'ok': True}


# ===== FAQ 导出 =====
@router.get('/faqs/export')
def export_faqs(
    format: str = 'xlsx',
    category: Optional[str] = None,
    published_only: bool = False,
    db: Session = Depends(get_db)
):
    """导出 FAQ 列表，支持 xlsx/csv/json/pdf/docx/md 格式"""
    q = db.query(FAQ)
    if published_only:
        q = q.filter(FAQ.is_published == True)
    if category:
        q = q.filter(FAQ.category == category)
    items = q.order_by(FAQ.created_at.desc()).all()

    faq_data = [{
        'id': f.id, 'question': f.question, 'answer': f.answer,
        'category': f.category or '', 'is_published': f.is_published,
    } for f in items]

    fmt = (format or 'xlsx').lower()
    now_str = datetime.now().strftime('%Y%m%d_%H%M%S')

    # --- JSON ---
    if fmt == 'json':
        import json
        from fastapi.responses import JSONResponse
        return JSONResponse(content=faq_data)

    # --- CSV ---
    if fmt == 'csv':
        import csv
        from io import StringIO
        buf = StringIO()
        writer = csv.DictWriter(buf, fieldnames=['id', 'question', 'answer', 'category', 'is_published'])
        writer.writeheader()
        writer.writerows(faq_data)
        from fastapi.responses import StreamingResponse
        return StreamingResponse(
            iter([buf.getvalue()]),
            media_type='text/csv',
            headers={'Content-Disposition': f'attachment; filename=faqs_{now_str}.csv'}
        )

    # --- Markdown ---
    if fmt == 'md':
        lines = ['# FAQ 常见问题解答', '']
        current_cat = None
        for item in faq_data:
            cat = item['category'] or '未分类'
            if cat != current_cat:
                lines.append(f'## {cat}')
                lines.append('')
                current_cat = cat
            published_tag = ' ✅' if item['is_published'] else ''
            lines.append(f'### Q{item["id"]}: {item["question"]}{published_tag}')
            lines.append('')
            lines.append(item['answer'] or '（暂无回答）')
            lines.append('')
        md_content = '
'.join(lines)
        from fastapi.responses import StreamingResponse
        return StreamingResponse(
            iter([md_content.encode('utf-8')]),
            media_type='text/markdown; charset=utf-8',
            headers={'Content-Disposition': f'attachment; filename=faqs_{now_str}.md'}
        )

    # --- Word (docx) ---
    if fmt == 'docx':
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches
        from fastapi.responses import StreamingResponse
        doc = DocxDocument()
        doc.add_heading('FAQ 常见问题解答', level=0)
        current_cat = None
        for item in faq_data:
            cat = item['category'] or '未分类'
            if cat != current_cat:
                doc.add_heading(cat, level=1)
                current_cat = cat
            p = doc.add_heading(f'Q{item["id"]}: {item["question"]}', level=2)
            doc.add_paragraph(item['answer'] or '（暂无回答）')
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename=faqs_{now_str}.docx'}
        )

    # --- PDF ---
    if fmt == 'pdf':
        try:
            import pdfkit
        except ImportError:
            raise HTTPException(500, 'pdfkit 未安装，无法导出 PDF')
        # Build HTML content
        html_lines = [
            '<html><head><meta charset="utf-8">',
            '<style>body{font-family:sans-serif;margin:40px}h1{color:#333}h2{color:#555;border-bottom:1px solid #ccc;padding-bottom:4px}h3{color:#666}.question{font-weight:bold;margin-top:16px}.answer{margin:8px 0 20px 20px;color:#444}</style>',
            '</head><body>',
            '<h1>FAQ 常见问题解答</h1>',
        ]
        current_cat = None
        for item in faq_data:
            cat = item['category'] or '未分类'
            if cat != current_cat:
                html_lines.append(f'<h2>{cat}</h2>')
                current_cat = cat
            published_tag = ' <span style="color:green">✅已发布</span>' if item['is_published'] else ''
            html_lines.append(f'<div class="question">Q{item["id"]}: {item["question"]}{published_tag}</div>')
            html_lines.append(f'<div class="answer">{(item["answer"] or "（暂无回答）")}</div>')
        html_lines.append('</body></html>')
        html_content = '
'.join(html_lines)

        pdf_buf = io.BytesIO()
        try:
            pdfkit.from_string(html_content, False, options={'encoding': 'utf-8', 'quiet': ''})
            # If pdfkit returns bytes
            from fastapi.responses import StreamingResponse
            pdf_bytes = pdfkit.from_string(html_content, False, options={'encoding': 'utf-8', 'quiet': ''})
            return StreamingResponse(
                io.BytesIO(pdf_bytes),
                media_type='application/pdf',
                headers={'Content-Disposition': f'attachment; filename=faqs_{now_str}.pdf'}
            )
        except Exception as e:
            # Fallback: try weasyprint
            try:
                from weasyprint import HTML as WeasyHTML
                pdf_bytes = WeasyHTML(string=html_content).write_pdf()
                from fastapi.responses import StreamingResponse
                return StreamingResponse(
                    io.BytesIO(pdf_bytes),
                    media_type='application/pdf',
                    headers={'Content-Disposition': f'attachment; filename=faqs_{now_str}.pdf'}
                )
            except ImportError:
                raise HTTPException(500, 'PDF 生成库(pdfkit/weasyprint)均未安装，无法导出 PDF')

    # --- Excel (default) ---
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = 'FAQ'
    ws.append(['ID', '问题', '回答', '分类', '是否发布'])
    for item in faq_data:
        ws.append([item['id'], item['question'], item['answer'], item['category'], '是' if item['is_published'] else '否'])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        buf,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename=faqs_{now_str}.xlsx'}
    )


# ===== 文书生成 =====
@router.get('/document-templates')
def list_templates(db: Session = Depends(get_db)):
    items = db.query(DocumentTemplate).all()
    return [{
        'id': t.id, 'name': t.name, 'template_type': t.template_type,
        'content': t.content,
    } for t in items]


@router.post('/document-templates')
def create_template(data: dict, db: Session = Depends(get_db)):
    t = DocumentTemplate(**data)
    db.add(t)
    db.commit()
    db.refresh(t)
    return {'id': t.id}


@router.post('/documents/generate')
def generate_document(data: dict, db: Session = Depends(get_db)):
    """基于模板和学生数据生成文书"""
    student_id = data.get('student_id')
    template_id = data.get('template_id')
    doc_type = data.get('doc_type', '')
    title = data.get('title', '')

    content = ''
    if template_id:
        tpl = db.query(DocumentTemplate).get(template_id)
        if tpl:
            content = tpl.content

    if student_id:
        stu = db.query(Student).get(student_id)
        if stu:
            cls_name = stu.class_obj.class_name if stu.class_obj else ''
            major_name = ''
            if stu.class_obj and stu.class_obj.major:
                major_name = stu.class_obj.major.major_name
            content = content.replace('{{姓名}}', stu.name or '')
            content = content.replace('{{学号}}', stu.student_no or '')
            content = content.replace('{{性别}}', stu.gender or '')
            content = content.replace('{{专业}}', major_name)
            content = content.replace('{{班级}}', cls_name)
            content = content.replace('{{政治面貌}}', stu.political_status or '群众')

    doc = GeneratedDocument(
        student_id=student_id, template_id=template_id,
        title=title, content=content, doc_type=doc_type
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return {'id': doc.id, 'content': content}


@router.get('/documents')
def list_documents(student_id: Optional[int] = None, db: Session = Depends(get_db)):
    q = db.query(GeneratedDocument)
    if student_id:
        q = q.filter(GeneratedDocument.student_id == student_id)
    items = q.order_by(GeneratedDocument.created_at.desc()).all()
    result = []
    for d in items:
        stu = db.query(Student).get(d.student_id) if d.student_id else None
        result.append({
            'id': d.id, 'title': d.title, 'doc_type': d.doc_type,
            'content': d.content, 'student_name': stu.name if stu else '',
            'created_at': str(d.created_at),
        })
    return result


@router.get('/documents/{doc_id}')
def get_document(doc_id: int, db: Session = Depends(get_db)):
    d = db.query(GeneratedDocument).get(doc_id)
    if not d:
        raise HTTPException(404, '文档不存在')
    stu = db.query(Student).get(d.student_id) if d.student_id else None
    return {
        'id': d.id, 'title': d.title, 'doc_type': d.doc_type,
        'content': d.content, 'student_name': stu.name if stu else '',
        'created_at': str(d.created_at),
    }


@router.put('/documents/{did}')
def update_document(did: int, data: dict, db: Session = Depends(get_db)):
    d = db.query(GeneratedDocument).get(did)
    if not d:
        raise HTTPException(404)
    for k, v in data.items():
        if hasattr(d, k):
            setattr(d, k, v)
    db.commit()
    return {'ok': True}


@router.delete('/documents/{did}')
def delete_document(did: int, db: Session = Depends(get_db)):
    d = db.query(GeneratedDocument).get(did)
    if d:
        db.delete(d)
        db.commit()
    return {'ok': True}


# ===== 周汇总 =====
@router.get('/weekly-summaries')
def list_weekly_summaries(db: Session = Depends(get_db)):
    items = db.query(WeeklySummary).order_by(WeeklySummary.created_at.desc()).all()
    from models import WeeklySummary as WS
    return [{
        'id': s.id, 'week_start': s.week_start, 'week_end': s.week_end,
        'content': s.content, 'summary_type': s.summary_type,
        'created_at': str(s.created_at),
    } for s in items]


@router.get('/weekly-summaries/{summary_id}')
def get_weekly_summary(summary_id: int, db: Session = Depends(get_db)):
    s = db.query(WeeklySummary).get(summary_id)
    if not s:
        raise HTTPException(404, '周汇总不存在')
    return {
        'id': s.id, 'week_start': s.week_start, 'week_end': s.week_end,
        'content': s.content, 'summary_type': s.summary_type,
        'created_at': str(s.created_at),
    }


@router.post('/weekly-summaries/generate')
def generate_weekly_summary(payload: dict = None, db: Session = Depends(get_db)):
    """自动生成周汇总（v3h-hotfix1：三 format 真差异化 + 日程本本周回顾/下周计划）
    payload 可选字段：
      dimensions: [academic|party|psychology|aid|employment|daily|activity] 多选
      format: bullet | paragraph | mixed
      week_offset: int，0=本周，-1=上周
    """
    from datetime import datetime, timedelta
    from models import (
        PsychologyRecord, FamilyContact, PartyProgress, WarningRecord,
        ClassMeeting, WeeklySummary as WS,
        Activity, StudentGrant, EmploymentRecord, GradeRecord,
        StudentDiscipline, StudentDormVisit, Student, ClassModel,
        Note, Countdown, Project
    )

    payload = payload or {}
    dims = payload.get('dimensions') or ['academic','party','psychology','aid','employment','daily','activity']
    fmt  = payload.get('format') or 'bullet'
    week_offset = int(payload.get('week_offset') or 0)

    today = datetime.now() + timedelta(weeks=week_offset)
    week_start = today - timedelta(days=today.weekday())
    week_end = week_start + timedelta(days=6)
    ws = week_start.strftime('%Y-%m-%d')
    we = week_end.strftime('%Y-%m-%d')
    # 下周
    nws = (week_start + timedelta(days=7)).strftime('%Y-%m-%d')
    nwe = (week_end + timedelta(days=7)).strftime('%Y-%m-%d')

    _stu_cache = {}
    def _stu(sid):
        if sid in _stu_cache:
            return _stu_cache[sid]
        s = db.query(Student).get(sid) if sid else None
        if s:
            cn = s.class_obj.class_name if s.class_obj else ''
            _stu_cache[sid] = (s.name or '未命名', cn)
        else:
            _stu_cache[sid] = ('未知学生', '')
        return _stu_cache[sid]

    def _fmt_date(d):
        if not d:
            return '日期未记录'
        try:
            dt = datetime.strptime(d[:10], '%Y-%m-%d')
            wk = ['一','二','三','四','五','六','日'][dt.weekday()]
            return f'{dt.strftime("%m-%d")}（周{wk}）'
        except Exception:
            return d

    def _in_range(d, a, b):
        if not d:
            return False
        try:
            return a <= d[:10] <= b
        except Exception:
            return False

    def _in_week(d):
        return _in_range(d, ws, we)
    def _in_next_week(d):
        return _in_range(d, nws, nwe)

    MAX_PER_DIM = 6

    # ---------- 通用渲染 ----------
    # section_dict 结构：
    #   title: '## 📚 学业跟踪'
    #   header: '本周新增学业预警 3 条：'（bullet/mixed 用）
    #   narrative: '本周新增红色预警 2 人（张三、李四），黄色预警 1 人……' （paragraph 主要输出）
    #   bullets: ['- 07-15（周一）· 🔴红牌 · 张三（软件2班）· 详见预警记录', ...]
    #   kpis: [('🔴', '红色', 2), ('🟡', '黄色', 1)]  # mixed 用大数字卡片
    #   tail: '**学业总览**：累计挂科 12 条，历史预警 30 条。'
    def _emoji_bar(kpis, width=20):
        """字符艺术水平条形图 · 每行一条"""
        if not kpis:
            return ''
        total = max(1, sum(k[2] for k in kpis) or 1)
        rows = []
        for emoji, label, count in kpis:
            n = int(round(count * width / total)) if total else 0
            bar = '█' * n + '░' * (width - n)
            rows.append(f'`{bar}` {emoji} {label} × **{count}**')
        return '\n'.join(rows)

    def _render(sec, fmt):
        title = sec.get('title', '')
        header = sec.get('header', '')
        narrative = sec.get('narrative', '')
        bullets = sec.get('bullets', [])[:MAX_PER_DIM]
        omitted = max(0, len(sec.get('bullets', [])) - MAX_PER_DIM)
        kpis = sec.get('kpis', [])
        tail = sec.get('tail', '')

        if fmt == 'paragraph':
            # 完全叙事段落
            body = title + '\n\n'
            body += (narrative or header) + '\n'
            if omitted:
                body += f'\n另有 {omitted} 条详情从略，可切换"要点列表"模式查看。\n'
            if tail:
                body += '\n' + tail + '\n'
            return body + '\n'

        if fmt == 'mixed':
            # 段落 + 大数字卡片 + 字符条形图 + top3 bullet
            body = title + '\n\n'
            body += (narrative or header) + '\n\n'
            # 大数字卡片 · emoji 版
            if kpis:
                cards = ' &nbsp; '.join([f'{k[0]} **{k[1]}** `{k[2]}`' for k in kpis])
                body += '> ' + cards + '\n\n'
                bar_txt = _emoji_bar(kpis)
                if bar_txt:
                    body += bar_txt + '\n\n'
            # top3 bullet
            if bullets:
                top = bullets[:3]
                body += '**代表条目：**\n' + '\n'.join(top) + '\n'
                if len(bullets) > 3:
                    rest = '；'.join([b[2:].split(' · ',1)[-1].strip() for b in bullets[3:]])
                    body += f'\n此外还包括：{rest}。\n'
            if omitted:
                body += f'\n（另有 {omitted} 条从略）\n'
            if tail:
                body += '\n' + tail + '\n'
            return body + '\n'

        # bullet（默认）· 完整列表
        body = title + '\n' + (header or narrative) + '\n'
        if bullets:
            body += '\n'.join(bullets) + '\n'
        if omitted:
            body += f'- （另有 {omitted} 条从略）\n'
        if tail:
            body += '\n' + tail + '\n'
        return body

    def _names(rs, key='student_id', max_n=8):
        """把一组 record 的学生姓名拼成'张三、李四、王五'"""
        seen = []
        for r in rs:
            sid = getattr(r, key, None) if not isinstance(r, dict) else r.get(key)
            name, _cn = _stu(sid)
            if name not in seen:
                seen.append(name)
            if len(seen) >= max_n:
                break
        s = '、'.join(seen)
        if len(rs) > len(seen):
            s += f' 等 {len(rs)} 人'
        return s

    sections = [f'# 第 {today.isocalendar()[1]} 周工作汇总  ·  {ws} ~ {we}\n',
                f'> 输出模式：**{ {"bullet":"要点列表","paragraph":"叙事段落","mixed":"图文混合"}.get(fmt, fmt) }**  ·  共 7 天\n']

    # ---------- 学业 ----------
    if 'academic' in dims:
        warns_all = db.query(WarningRecord).order_by(WarningRecord.created_at.desc()).all()
        warns_week = [w for w in warns_all if _in_week(str(w.created_at))]
        red = [w for w in warns_week if w.warning_type == 'red']
        yellow = [w for w in warns_week if w.warning_type != 'red']

        bullets = []
        for w in warns_week:
            name, cn = _stu(w.student_id)
            tag = '🔴红牌' if w.warning_type == 'red' else '🟡黄牌'
            desc = (w.description or '详见预警记录').strip()[:40]
            bullets.append(f'- {_fmt_date(str(w.created_at)[:10])} · {tag} · {name}（{cn}）· {desc}')

        # 叙事总结
        if warns_week:
            parts = []
            if red:
                parts.append(f'新增红色预警 {len(red)} 人（{_names(red)}），已按流程提醒到位')
            if yellow:
                parts.append(f'黄色预警 {len(yellow)} 人（{_names(yellow)}），持续跟进中')
            narrative = f'**学业维度**：本周' + '；'.join(parts) + '。'
        else:
            narrative = '**学业维度**：本周未新增学业预警，学业风险总体平稳。'

        fail_this = db.query(GradeRecord).filter(GradeRecord.score < 60).count()
        tail = f'**学业总览**：累计挂科记录 {fail_this} 条，历史预警 {len(warns_all)} 条。'
        kpis = []
        if red: kpis.append(('🔴','红色预警',len(red)))
        if yellow: kpis.append(('🟡','黄色预警',len(yellow)))

        sections.append(_render({
            'title': '## 📚 学业跟踪',
            'header': f'本周新增学业预警 **{len(warns_week)}** 条：',
            'narrative': narrative,
            'bullets': bullets, 'kpis': kpis, 'tail': tail,
        }, fmt))

    # ---------- 谈心谈话 ----------
    if 'psychology' in dims:
        recs = db.query(PsychologyRecord).all()
        week_recs = [r for r in recs if _in_week(r.record_date)]
        bullets = []
        for r in week_recs:
            name, cn = _stu(r.student_id)
            loc = r.location or '未记录地点'
            topic = r.topic or '常规谈话'
            summary = (r.summary or '').replace('\n',' ').strip()[:60]
            follow = f'（后续 {r.next_follow_date} 再跟进）' if r.next_follow_date else ''
            bullets.append(f'- {_fmt_date(r.record_date)} · 与 **{name}**（{cn}）在 {loc} 谈心 · 主题：{topic}' +
                           (f' · 要点：{summary}' if summary else '') + follow)
        # 关注等级分布
        lv_map = {'一': 0, '二': 0, '三': 0, '普通': 0}
        for r in week_recs:
            for k in lv_map:
                if k in (r.attention_level or ''):
                    lv_map[k] += 1
                    break

        if week_recs:
            narrative = f'**谈心谈话**：本周共开展 {len(week_recs)} 次心理关怀谈话，涉及 {_names(week_recs)}。'
            details = []
            if lv_map['一']: details.append(f'一级重点 {lv_map["一"]} 人')
            if lv_map['二']: details.append(f'二级关注 {lv_map["二"]} 人')
            if lv_map['三']: details.append(f'三级关注 {lv_map["三"]} 人')
            if details:
                narrative += ' 关注等级分布：' + '，'.join(details) + '。'
        else:
            narrative = '**谈心谈话**：本周未记录谈心谈话，建议主动约谈有情绪波动的学生。'

        kpis = []
        for lv, cnt in lv_map.items():
            if cnt:
                emoji = {'一':'🔴','二':'🟠','三':'🟢','普通':'⚪'}[lv]
                kpis.append((emoji, f'{lv}级', cnt))

        sections.append(_render({
            'title': '## 💚 谈心谈话',
            'header': f'本周开展谈心谈话 **{len(week_recs)}** 次：',
            'narrative': narrative,
            'bullets': bullets, 'kpis': kpis,
        }, fmt))

    # ---------- 党团 ----------
    if 'party' in dims:
        recs = db.query(PartyProgress).all()
        week_recs = [r for r in recs if _in_week(r.stage_date) or _in_week(str(r.created_at))]
        bullets = []
        stage_map = {}
        for r in week_recs:
            name, cn = _stu(r.student_id)
            notes = (r.notes or '').strip()[:40]
            stage_map[r.stage or '未标注'] = stage_map.get(r.stage or '未标注', 0) + 1
            bullets.append(f'- {_fmt_date(r.stage_date or str(r.created_at)[:10])} · {name}（{cn}）· 进入「**{r.stage}**」阶段' +
                           (f' · 联系人：{r.contact_person}' if r.contact_person else '') +
                           (f' · {notes}' if notes else ''))

        if week_recs:
            stage_parts = '；'.join([f'{k} {v} 人' for k, v in stage_map.items()])
            narrative = f'**党团发展**：本周共 {len(week_recs)} 项节点变动 —— {stage_parts}。涉及学生：{_names(week_recs)}。'
        else:
            narrative = '**党团发展**：本周无党团发展节点变动。'
        kpis = [('🚩', k, v) for k, v in stage_map.items()]

        sections.append(_render({
            'title': '## 🚩 党团建设',
            'header': f'本周党团发展节点 **{len(week_recs)}** 项：',
            'narrative': narrative,
            'bullets': bullets, 'kpis': kpis,
        }, fmt))

    # ---------- 资助 ----------
    if 'aid' in dims:
        recs = db.query(StudentGrant).all()
        week_recs = [r for r in recs if _in_week(str(r.created_at))]
        bullets = []
        total_amt = 0.0
        for r in week_recs:
            name, cn = _stu(r.student_id)
            total_amt += (r.amount or 0)
            bullets.append(f'- {_fmt_date(str(r.created_at)[:10])} · {name}（{cn}） · {r.grant_type or "资助"} ¥{r.amount:.0f}' +
                           (f' · {r.notes}' if r.notes else ''))
        if week_recs:
            narrative = f'**资助帮扶**：本周新增资助/助学金 {len(week_recs)} 条，涉及 {_names(week_recs)}，合计发放金额 ¥{total_amt:.0f}。'
        else:
            narrative = '**资助帮扶**：本周资助无变动，如临学期节点可发起复核。'
        kpis = [('💰','资助笔数',len(week_recs))] if week_recs else []

        sections.append(_render({
            'title': '## 💰 资助帮扶',
            'header': f'本周资助/助学金动态 **{len(week_recs)}** 条：',
            'narrative': narrative,
            'bullets': bullets, 'kpis': kpis,
        }, fmt))

    # ---------- 就业 ----------
    if 'employment' in dims:
        recs = db.query(EmploymentRecord).all()
        week_recs = [r for r in recs if _in_week(r.offer_date) or _in_week(str(r.created_at))]
        bullets = []
        status_map = {}
        for r in week_recs:
            name, cn = _stu(r.student_id)
            status_map[r.status or '未标注'] = status_map.get(r.status or '未标注', 0) + 1
            bits = []
            if r.status: bits.append(r.status)
            if r.target_industry: bits.append(r.target_industry)
            if r.target_position: bits.append(r.target_position)
            if r.internship_company: bits.append(f'实习→{r.internship_company}')
            if r.salary_range: bits.append(f'薪资 {r.salary_range}')
            bullets.append(f'- {_fmt_date(r.offer_date or str(r.created_at)[:10])} · {name}（{cn}）· ' + ' / '.join(bits))
        if week_recs:
            stat_parts = '；'.join([f'{k} {v} 人' for k, v in status_map.items()])
            narrative = f'**就业跟踪**：本周就业进展 {len(week_recs)} 条，状态分布：{stat_parts}。涉及学生：{_names(week_recs)}。'
        else:
            narrative = '**就业跟踪**：本周就业无进展，可推送校招信息或组织宣讲会。'
        kpis = [('🎯', k, v) for k, v in status_map.items()]

        sections.append(_render({
            'title': '## 🎯 就业跟踪',
            'header': f'本周就业进展 **{len(week_recs)}** 条：',
            'narrative': narrative,
            'bullets': bullets, 'kpis': kpis,
        }, fmt))

    # ---------- 日常（家校+违纪+走访） ----------
    if 'daily' in dims:
        fc = [r for r in db.query(FamilyContact).all() if _in_week(r.contact_date)]
        dc = [r for r in db.query(StudentDiscipline).all() if _in_week(r.discipline_date)]
        vs = [r for r in db.query(StudentDormVisit).all() if _in_week(r.visit_date)]

        bullets = []
        for r in fc:
            name, cn = _stu(r.student_id)
            bullets.append(f'- {_fmt_date(r.contact_date)} · 联系 {name}（{cn}）家长 {r.parent_name or ""} · 方式：{r.contact_method or "电话"} · 主题：{r.topic or "常规沟通"}')
        for r in dc:
            name, cn = _stu(r.student_id)
            bullets.append(f'- {_fmt_date(r.discipline_date)} · {name}（{cn}）· {r.level or ""}{r.discipline_type or "处分"} · 原因：{(r.reason or "").strip()[:40]}')
        for r in vs:
            name, cn = _stu(r.student_id)
            sit = (r.situation or '').strip()[:40]
            bullets.append(f'- {_fmt_date(r.visit_date)} · 走访 {name}（{cn}）· 寝室 {r.dorm_room or "?"} · 走访人：{r.visitor or "本人"}' + (f' · {sit}' if sit else ''))

        parts = []
        if fc: parts.append(f'家校沟通 {len(fc)} 次（{_names(fc)}）')
        if dc: parts.append(f'违纪处理 {len(dc)} 次（{_names(dc)}）')
        if vs: parts.append(f'宿舍走访 {len(vs)} 次（{_names(vs)}）')
        if parts:
            narrative = f'**日常事务**：本周' + '、'.join(parts) + '。'
        else:
            narrative = '**日常事务**：本周日常事务平稳，无违纪、无重要家校沟通、无宿舍走访记录。'
        kpis = []
        if fc: kpis.append(('📞','家校沟通',len(fc)))
        if dc: kpis.append(('⚠️','违纪',len(dc)))
        if vs: kpis.append(('🏠','宿舍走访',len(vs)))

        sections.append(_render({
            'title': '## 📖 日常事务',
            'header': f'本周日常事务合计 **{len(fc)+len(dc)+len(vs)}** 项：',
            'narrative': narrative,
            'bullets': bullets, 'kpis': kpis,
        }, fmt))

    # ---------- 活动（含班会） ----------
    if 'activity' in dims:
        acts = [a for a in db.query(Activity).all() if _in_week(a.activity_date)]
        meets = [m for m in db.query(ClassMeeting).all() if _in_week(m.meeting_date)]

        bullets = []
        for a in acts:
            bullets.append(f'- {_fmt_date(a.activity_date)} · **{a.title}**' +
                           (f' · 地点：{a.location}' if a.location else '') +
                           (f' · 类型：{a.activity_type}' if a.activity_type else '') +
                           (f' · 状态：{a.status}' if a.status else ''))
        for m in meets:
            cls = db.query(ClassModel).get(m.class_id) if m.class_id else None
            cn = cls.class_name if cls else '(未指定班级)'
            bullets.append(f'- {_fmt_date(m.meeting_date)} · {cn} · 主题：{m.topic or "常规班会"} · 出勤 {m.attendance_count or 0} 人' +
                           (f' · 结论：{(m.resolution or "").strip()[:40]}' if m.resolution else ''))

        parts = []
        if acts: parts.append(f'组织学院/学生活动 {len(acts)} 场（{"、".join([a.title for a in acts[:5]])}）')
        if meets: parts.append(f'开展班会 {len(meets)} 次')
        if parts:
            narrative = f'**活动 & 班会**：本周' + '，'.join(parts) + '。'
        else:
            narrative = '**活动 & 班会**：本周无组织的学院活动或班会，可结合校历安排下周事项。'
        kpis = []
        if acts: kpis.append(('🎨','活动',len(acts)))
        if meets: kpis.append(('🏫','班会',len(meets)))

        sections.append(_render({
            'title': '## 🎨 活动 & 班会',
            'header': f'本周活动 **{len(acts)}** 场 · 班会 **{len(meets)}** 次：',
            'narrative': narrative,
            'bullets': bullets, 'kpis': kpis,
        }, fmt))

    # ---------- 本周日程回顾（结合日程本：Note done + Countdown + Project 进度） ----------
    notes_done_week = [n for n in db.query(Note).all()
                       if n.status == 'done' and (_in_week(n.due_date) or _in_week(str(n.updated_at)))]
    countdowns_week = [c for c in db.query(Countdown).all() if _in_week(c.target_date)]
    projects_active = db.query(Project).filter(Project.status == 'active').all()
    projects_touched_week = [p for p in projects_active
                             if _in_week(p.start_date) or _in_week(p.end_date)]

    bullets_done = []
    for n in notes_done_week[:MAX_PER_DIM]:
        bullets_done.append(f'- ✅ {_fmt_date(n.due_date or str(n.updated_at)[:10])} · {n.title or "(无标题)"}' +
                            (f' · {(n.content or "").strip()[:30]}' if n.content else ''))
    for c in countdowns_week[:MAX_PER_DIM]:
        bullets_done.append(f'- 📅 {_fmt_date(c.target_date)} · **{c.title}**（{c.category or "校历"}）')
    for p in projects_touched_week[:MAX_PER_DIM]:
        bullets_done.append(f'- 🎯 项目 **{p.name}** · 进度 {p.progress}%' +
                            (f' · 起止 {p.start_date} ~ {p.end_date}' if p.start_date else ''))

    if bullets_done:
        narrative_done = (
            f'**本周日程回顾**：完成待办 {len(notes_done_week)} 项，'
            f'触达校历节点 {len(countdowns_week)} 个，推进中的专项工作 {len(projects_touched_week)} 项。'
        )
    else:
        narrative_done = '**本周日程回顾**：本周日程本无已完成待办或校历节点，可按计划稳步推进。'
    sections.append(_render({
        'title': '## 🗓️ 本周日程回顾',
        'header': f'本周完成 {len(notes_done_week)} 项待办 · 校历节点 {len(countdowns_week)} 个 · 项目触达 {len(projects_touched_week)} 项',
        'narrative': narrative_done,
        'bullets': bullets_done,
        'kpis': [('✅','已完成待办',len(notes_done_week)),
                 ('📅','校历节点',len(countdowns_week)),
                 ('🎯','项目触达',len(projects_touched_week))],
    }, fmt))

    # ---------- 下周计划提醒 ----------
    notes_todo_next = [n for n in db.query(Note).all()
                       if n.category == 'todo' and n.status != 'done' and _in_next_week(n.due_date)]
    countdowns_next = [c for c in db.query(Countdown).all() if _in_next_week(c.target_date)]
    acts_next = [a for a in db.query(Activity).all() if _in_next_week(a.activity_date)]
    meets_next = [m for m in db.query(ClassMeeting).all() if _in_next_week(m.meeting_date)]

    bullets_next = []
    for n in notes_todo_next[:MAX_PER_DIM]:
        prio = ['低','中','高'][n.priority or 0] if (n.priority or 0) <= 2 else '高'
        bullets_next.append(f'- ☑️ {_fmt_date(n.due_date)} · [{prio}] {n.title or "(无标题)"}')
    for c in countdowns_next[:MAX_PER_DIM]:
        bullets_next.append(f'- 📅 {_fmt_date(c.target_date)} · **{c.title}**（{c.category or "校历"}）')
    for a in acts_next[:MAX_PER_DIM]:
        bullets_next.append(f'- 🎨 {_fmt_date(a.activity_date)} · {a.title}' + (f' · {a.location}' if a.location else ''))
    for m in meets_next[:MAX_PER_DIM]:
        cls = db.query(ClassModel).get(m.class_id) if m.class_id else None
        cn = cls.class_name if cls else '(未指定班级)'
        bullets_next.append(f'- 🏫 {_fmt_date(m.meeting_date)} · {cn} · {m.topic or "班会"}')

    total_next = len(notes_todo_next) + len(countdowns_next) + len(acts_next) + len(meets_next)
    if total_next:
        narrative_next = (
            f'**下周计划提醒**（{nws} ~ {nwe}）：待办 {len(notes_todo_next)} 项、校历节点 {len(countdowns_next)} 个、'
            f'活动 {len(acts_next)} 场、班会 {len(meets_next)} 次，请提前安排时间与资源。'
        )
    else:
        narrative_next = f'**下周计划提醒**（{nws} ~ {nwe}）：日程本暂无预设事项，可尽早规划下周重点工作。'
    sections.append(_render({
        'title': '## 📌 下周计划提醒',
        'header': f'下周共 {total_next} 项预设事项：',
        'narrative': narrative_next,
        'bullets': bullets_next,
        'kpis': [('☑️','待办',len(notes_todo_next)),
                 ('📅','校历',len(countdowns_next)),
                 ('🎨','活动',len(acts_next)),
                 ('🏫','班会',len(meets_next))],
    }, fmt))

    # ---------- 收尾 ----------
    sections.append('---\n')
    sections.append('*本汇总由系统按数据库真实记录自动汇编 · 三种模式（要点列表 / 叙事段落 / 图文混合）根据 format 参数切换 · 结尾已联动日程本给出本周回顾与下周计划。*')

    content = '\n'.join(sections)

    summary = WS(week_start=ws, week_end=we, content=content, summary_type='auto',
                 title=f'第{today.isocalendar()[1]}周工作汇总（{ws}~{we}）')
    db.add(summary)
    db.commit()
    db.refresh(summary)
    return {'id': summary.id, 'content': content, 'dimensions': dims, 'format': fmt}


@router.put('/weekly-summaries/{sid}')
def update_weekly_summary(sid: int, data: dict, db: Session = Depends(get_db)):
    from models import WeeklySummary as WS
    s = db.query(WS).get(sid)
    if not s:
        raise HTTPException(404)
    for k, v in data.items():
        if hasattr(s, k):
            setattr(s, k, v)
    db.commit()
    return {'ok': True}


@router.delete('/weekly-summaries/{sid}')
def delete_weekly_summary(sid: int, db: Session = Depends(get_db)):
    from models import WeeklySummary as WS
    s = db.query(WS).get(sid)
    if s:
        db.delete(s)
        db.commit()
    return {'ok': True}


# ===== AI 智能导入 =====
@router.post('/smart-import/preview')
def smart_import_preview(data: dict, db: Session = Depends(get_db)):
    """AI智能导入预览 - 解析列名并映射"""
    import re
    rows = data.get('rows', [])
    headers = data.get('headers', [])

    if not headers or not rows:
        return {'error': '无数据'}

    # 同义词映射表
    synonym_map = {
        'student_no': ['学号', '学生编号', '学生号', '编号', 'student_no', 'student_id'],
        'name': ['姓名', '学生姓名', '名字', '名', 'name'],
        'gender': ['性别', '男女', 'sex', 'gender'],
        'class_name': ['班级', '所在班级', '班级号', 'class', 'class_name'],
        'major': ['专业', '所学专业', '专业名称', 'major'],
        'political_status': ['政治面貌', '政治身份', 'political_status'],
        'phone': ['联系电话', '手机号', '电话', 'phone', 'tel'],
        'parent_phone': ['家长电话', '紧急联系电话', '家庭联系方式', 'parent_phone'],
        'birth_date': ['出生日期', '出生年月', 'birthday', 'birth_date'],
        'birth_source': ['生源地', '籍贯', '生源省份', 'birth_source'],
        'email': ['邮箱', '电子邮件', 'email'],
        'family_situation': ['家庭情况', '家庭状况', 'family_situation'],
    }

    # 自动映射
    mapping = {}
    for i, h in enumerate(headers):
        h_clean = h.strip()
        for field, synonyms in synonym_map.items():
            if h_clean in synonyms or any(s in h_clean for s in synonyms):
                mapping[i] = {'header': h_clean, 'field': field}
                break
        if i not in mapping:
            mapping[i] = {'header': h_clean, 'field': ''}

    # 前5行预览
    preview_rows = []
    for row in rows[:5]:
        preview_rows.append([str(cell) if cell is not None else '' for cell in row])

    return {'mapping': mapping, 'preview_rows': preview_rows, 'total_rows': len(rows)}


@router.post('/smart-import/execute')
def smart_import_execute(data: dict, db: Session = Depends(get_db)):
    """执行智能导入"""
    from models import Student
    rows = data.get('rows', [])
    mapping = data.get('mapping', {})
    conflict_action = data.get('conflict_action', 'skip')  # skip/overwrite

    # 反转映射: field -> column_index
    field_to_col = {}
    for col_idx_str, m in mapping.items():
        col_idx = int(col_idx_str)
        field = m.get('field', '')
        if field:
            field_to_col[field] = col_idx

    imported = 0
    skipped = 0
    for row in rows:
        student_data = {}
        for field, col_idx in field_to_col.items():
            if col_idx < len(row):
                val = row[col_idx]
                student_data[field] = str(val).strip() if val is not None else ''

        student_no = student_data.get('student_no', '')
        if not student_no:
            skipped += 1
            continue

        existing = db.query(Student).filter(Student.student_no == student_no).first()
        if existing:
            if conflict_action == 'skip':
                skipped += 1
                continue
            elif conflict_action == 'overwrite':
                for k, v in student_data.items():
                    if k != 'student_no':
                        setattr(existing, k, v)
                imported += 1
            else:  # keep_both
                s = Student(**student_data)
                db.add(s)
                imported += 1
        else:
            s = Student(**student_data)
            db.add(s)
            imported += 1

    db.commit()
    return {'imported': imported, 'skipped': skipped, 'total': len(rows)}
